"""
Contract generator — bouwt zowel een rendered HTML (string) als een DOCX
(file) voor een gegeven klant.

CLI:
  python3 sidequest/generate_contract_docx.py
  → genereert sidequest/Contract_DeCargoWinkel.docx met de hardcoded
    De-Cargo-Winkel data (handig als fallback / backup).

Module-import:
  from sidequest.generate_contract_docx import build_contract_docx, render_contract_html
  build_contract_docx(client_data, output_path)
  render_contract_html(client_data, docx_url='...') -> str
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colors ───────────────────────────────────────────────────────────────
INK     = RGBColor(0x0A, 0x06, 0x18)
INK_2   = RGBColor(0x2a, 0x24, 0x40)
MUTED   = RGBColor(0x5a, 0x54, 0x70)
ACCENT  = RGBColor(0x1D, 0x4E, 0xD8)

# ── Default De Cargo Winkel client (fallback / CLI) ──────────────────────
DEFAULT_CLIENT = {
    'CLIENT_BEDRIJF':       'De Cargo Winkel',
    'CLIENT_NAAM_TEKENAAR': 'Theo Baan',
    'CLIENT_ADRES':         'Anthonie Fokkerstraat 61 N, 3772 MP Barneveld',
    'CLIENT_KVK':           '__________________',
    'CLIENT_BTW':           '__________________',
    'CLIENT_EMAIL':         '__________________',
}

# ── Token mapping: onboarding-row keys → contract placeholders ───────────
def _from_onboarding_row(row):
    """Map a row from the onboarding table (dict) to contract tokens.
    Empty values fall back to '__________________' so the printed contract
    still has visible blanks for fields the prospect didn't fill in."""
    blank = '__________________'
    def v(key):
        x = (row or {}).get(key)
        return x.strip() if isinstance(x, str) and x.strip() else blank
    return {
        'CLIENT_BEDRIJF':       v('naam'),
        'CLIENT_NAAM_TEKENAAR': v('contact'),
        'CLIENT_ADRES':         v('adres'),
        'CLIENT_KVK':           v('kvk'),
        'CLIENT_BTW':           v('btw'),
        'CLIENT_EMAIL':         v('email'),
    }

# ── Helpers for DOCX building ────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def add_run(paragraph, text, bold=False, size=None, color=None, italic=False):
    run = paragraph.add_run(text)
    run.font.name = 'Calibri'
    if bold:    run.bold = True
    if italic:  run.italic = True
    if size:    run.font.size = Pt(size)
    if color:   run.font.color.rgb = color
    return run

def small(doc, text, size=9, color=MUTED, italic=False, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    add_run(p, text, size=size, color=color, italic=italic)
    return p

def article_heading(doc, num, name):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14); pf.space_after = Pt(4); pf.keep_with_next = True
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4');    bottom.set(qn('w:color'), 'd8d3e3')
    pBdr.append(bottom); pPr.append(pBdr)
    add_run(p, f'Art. {num}  ', bold=True, size=10, color=ACCENT)
    add_run(p, name, bold=True, size=12, color=INK)

def clause(doc, num, text, indent_cm=0.6):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(indent_cm); pf.space_after = Pt(4)
    add_run(p, num + ' ', bold=True, color=INK_2, size=10.5)
    add_run(p, text, color=INK, size=10.5)

def bullet(doc, text, indent_cm=1.3):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(indent_cm); p.paragraph_format.space_after = Pt(2)
    add_run(p, text, color=INK, size=10.5)

def callout(doc, text):
    tbl = doc.add_table(rows=1, cols=1); tbl.autofit = False
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, 'FFF7E6'); cell.width = Cm(15.8)
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '18'); left.set(qn('w:color'), 'D97706')
    borders.append(left); tcPr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    add_run(p, text, color=INK_2, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def signature_block(doc, client):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24); p.paragraph_format.space_after = Pt(8)
    add_run(p, 'Aldus overeengekomen en in tweevoud ondertekend', bold=True, size=13, color=INK)
    tbl = doc.add_table(rows=1, cols=2); tbl.autofit = False
    parties = [
        {'title': 'OPDRACHTNEMER', 'name': 'Julian Verboom',
         'co': 'handelend voor Viral Conversions'},
        {'title': 'OPDRACHTGEVER', 'name': client['CLIENT_NAAM_TEKENAAR'],
         'co': f"handelend voor {client['CLIENT_BEDRIJF']}"},
    ]
    for col_idx, party in enumerate(parties):
        c = tbl.rows[0].cells[col_idx]; c.width = Cm(7.9)
        p1 = c.paragraphs[0]; add_run(p1, party['title'], bold=True, size=9, color=ACCENT)
        p2 = c.add_paragraph(); p2.paragraph_format.space_before = Pt(2)
        add_run(p2, party['name'], bold=True, size=11, color=INK)
        p3 = c.add_paragraph(); add_run(p3, party['co'], size=9, color=MUTED, italic=True)
        sp = c.add_paragraph(); sp.paragraph_format.space_before = Pt(48)
        sline = c.add_paragraph()
        sPr = sline._p.get_or_add_pPr(); sBdr = OxmlElement('w:pBdr')
        sBot = OxmlElement('w:bottom')
        sBot.set(qn('w:val'), 'single'); sBot.set(qn('w:sz'), '8'); sBot.set(qn('w:color'), '000000')
        sBdr.append(sBot); sPr.append(sBdr)
        for k in ('Plaats', 'Datum', 'Handtekening'):
            m = c.add_paragraph(); m.paragraph_format.space_after = Pt(2)
            add_run(m, f'{k}: ', bold=True, size=9, color=MUTED)
            add_run(m, '______________________________', size=9, color=INK)

# ── Build a complete DOCX for the given client (dict of placeholders) ────
def build_contract_docx(client, output_path):
    """Generate the DOCX file for one client. `client` is a dict of CLIENT_* tokens."""
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(10.5); style.font.color.rgb = INK

    # Header
    p = doc.add_paragraph()
    add_run(p, 'VIRAL CONVERSIONS', bold=True, size=14, color=INK)
    small(doc, 'OVEREENKOMST VAN OPDRACHT', size=9, color=MUTED)
    small(doc, f'Document: Contract — Webshop {client["CLIENT_BEDRIJF"]}', size=9, color=MUTED)
    small(doc, 'Datum: __________________', size=9, color=MUTED, space_after=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    add_run(p, 'Overeenkomst van opdracht', bold=True, size=20, color=INK)
    small(doc, f'Voor de bouw, oplevering en het onderhoud van een webshop voor {client["CLIENT_BEDRIJF"]}.',
          size=10.5, color=MUTED, space_after=14)

    # Parties block
    parties_tbl = doc.add_table(rows=1, cols=2); parties_tbl.autofit = False
    parties_data = [
        {
            'title': 'OPDRACHTNEMER',
            'name': 'Julian Verboom',
            'sub':  'handelend onder de naam Viral Conversions',
            'rows': [
                ('Adres',  'Koolwijkseweg 25a, Stolwijk'),
                ('KvK',    '99922533'),
                ('Btw-id', 'NL267949303B01'),
                ('IBAN',   'NL90 RABO 0172 1492 82 (t.n.v. Viral Conversions)'),
            ],
        },
        {
            'title': 'OPDRACHTGEVER',
            'name': client['CLIENT_NAAM_TEKENAAR'],
            'sub':  f"handelend onder de naam {client['CLIENT_BEDRIJF']}",
            'rows': [
                ('Adres',  client['CLIENT_ADRES']),
                ('KvK',    client['CLIENT_KVK']),
                ('Btw-id', client['CLIENT_BTW']),
                ('E-mail', client['CLIENT_EMAIL']),
            ],
        },
    ]
    for col_idx, party in enumerate(parties_data):
        cell = parties_tbl.rows[0].cells[col_idx]; cell.width = Cm(7.9)
        p = cell.paragraphs[0]; add_run(p, party['title'], bold=True, size=9, color=ACCENT)
        pn = cell.add_paragraph(); add_run(pn, party['name'], bold=True, size=13, color=INK)
        ps = cell.add_paragraph(); add_run(ps, party['sub'], size=10, color=INK_2, italic=True)
        for lbl, val in party['rows']:
            prow = cell.add_paragraph(); prow.paragraph_format.space_after = Pt(1)
            add_run(prow, f'{lbl}: ', size=9.5, color=MUTED)
            add_run(prow, val, size=9.5, color=INK_2)
    small(doc, '', space_after=8)

    # In aanmerking nemende
    cons_tbl = doc.add_table(rows=1, cols=1); cons_tbl.autofit = False
    ccell = cons_tbl.rows[0].cells[0]
    set_cell_bg(ccell, 'EFF4FE')
    tcPr = ccell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders'); lb = OxmlElement('w:left')
    lb.set(qn('w:val'), 'single'); lb.set(qn('w:sz'), '18'); lb.set(qn('w:color'), '1D4ED8')
    borders.append(lb); tcPr.append(borders); ccell.width = Cm(15.8)
    p = ccell.paragraphs[0]; add_run(p, 'In aanmerking nemende dat:', bold=True, size=9.5, color=INK_2)
    ccell.add_paragraph().add_run(
        'a. de opdrachtgever een webshop wenst voor de verkoop van producten via internet, met een eigen '
        'beheeromgeving waarmee de opdrachtgever zelfstandig producten, prijzen, voorraad en bestellingen '
        'kan beheren;'
    ).font.size = Pt(9.5)
    ccell.add_paragraph().add_run(
        'b. de opdrachtnemer als professional in het ontwerp en de bouw van webshops bereid is deze opdracht '
        'onder de hierna omschreven voorwaarden te aanvaarden;'
    ).font.size = Pt(9.5)
    ccell.add_paragraph().add_run('komen partijen het volgende overeen.').font.size = Pt(9.5)
    small(doc, '', space_after=6)

    # Article 1
    article_heading(doc, 1, 'Partijen en definities')
    clause(doc, '1.1', 'Onder "opdrachtnemer" wordt verstaan: Julian Verboom, handelend onder de naam Viral Conversions, zoals hierboven aangeduid.')
    clause(doc, '1.2', f'Onder "opdrachtgever" wordt verstaan: {client["CLIENT_NAAM_TEKENAAR"]}, handelend onder de naam {client["CLIENT_BEDRIJF"]}, zoals hierboven aangeduid.')
    clause(doc, '1.3', 'Onder "de webshop" wordt verstaan: het volledige in artikel 2 omschreven softwareproduct, inclusief beheeromgeving en betaalkoppeling.')
    clause(doc, '1.4', 'Onder "scope" wordt verstaan: het geheel van expliciet in deze overeenkomst opgenomen werkzaamheden en deliverables.')
    clause(doc, '1.5', 'Onder "meerwerk" wordt verstaan: alle werkzaamheden die buiten de in artikel 2 vastgelegde scope vallen.')

    # Article 2
    article_heading(doc, 2, 'De opdracht en scope')
    clause(doc, '2.1', f'De opdrachtgever verleent aan de opdrachtnemer de opdracht tot het ontwerp, de bouw en de oplevering van een complete webshop voor {client["CLIENT_BEDRIJF"]}.')
    clause(doc, '2.2', 'Inbegrepen in de bouwprijs als bedoeld in artikel 3 zijn:')
    bullet(doc, 'Een complete webshop, bestaande uit productpagina’s, een winkelwagen en een complete checkout-flow.')
    bullet(doc, 'Een veilige betaalkoppeling via Mollie of Stripe, naar keuze van de opdrachtgever. Het betaalproviders-account staat op naam van de opdrachtgever en wordt door of namens de opdrachtgever aangevraagd.')
    bullet(doc, 'Een beheerdashboard waarin de opdrachtgever zelfstandig producten kan toevoegen en aanpassen, voorraad kan beheren, prijzen kan wijzigen en bestellingen kan inzien.')
    bullet(doc, 'Beveiliging conform gangbare standaarden, waaronder een geldig SSL-certificaat.')
    bullet(doc, 'AVG-conformiteit van de webshop-fundamenten (cookiebanner, dataminimalisatie, plaats voor privacyverklaring en algemene voorwaarden).')
    clause(doc, '2.3', 'Niet inbegrepen in de bouwprijs en uitdrukkelijk buiten de scope:')
    bullet(doc, 'Het schrijven of leveren van algemene voorwaarden, privacyverklaring, cookieverklaring en andere juridische teksten.')
    bullet(doc, 'Het aanleveren van productcontent (teksten, afbeeldingen, prijzen, varianten, voorraad).')
    bullet(doc, 'Het importeren van bestaande productdata uit een ander systeem indien dit een niet-standaardformaat betreft; zie artikel 6.')
    bullet(doc, 'Marketing, SEO-trajecten, advertentiecampagnes, fotografie en copywriting.')
    bullet(doc, 'Integraties met externe systemen (boekhouding, ERP, koeriers-API’s, marktplaatsen) anders dan de in artikel 2.2 genoemde betaalkoppeling.')
    bullet(doc, 'Alles wat niet expliciet in artikel 2.2 is opgenomen.')
    clause(doc, '2.4', 'Werkzaamheden buiten de scope zijn meerwerk en worden uitgevoerd conform artikel 6.')

    # Article 3 (payment table)
    article_heading(doc, 3, 'Prijs en betaling')
    clause(doc, '3.1', 'De totale prijs voor de bouw van de webshop zoals omschreven in artikel 2.2 bedraagt € 800,00 exclusief btw. Tenzij anders schriftelijk afgesproken is dit een vaste prijs voor de overeengekomen scope.')
    clause(doc, '3.2', 'Het bedrag wordt in twee termijnen voldaan:')
    pay_tbl = doc.add_table(rows=4, cols=4); pay_tbl.style = 'Light Grid Accent 1'
    pay_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = pay_tbl.rows[0].cells
    for i, txt in enumerate(['Termijn', 'Moment van facturering', 'Excl. btw', 'Incl. btw (21%)']):
        set_cell_bg(hdr[i], 'F5F3FA')
        add_run(hdr[i].paragraphs[0], txt, bold=True, size=10, color=INK_2)
    rows_data = [
        ('40%',    'Aanbetaling. De bouw start na ontvangst van deze aanbetaling.',           '€ 320,00', '€ 387,20'),
        ('60%',    'Na oplevering en schriftelijke goedkeuring door de opdrachtgever.',       '€ 480,00', '€ 580,80'),
        ('Totaal', '',                                                                          '€ 800,00', '€ 968,00'),
    ]
    for r_idx, (a, b, c_val, d) in enumerate(rows_data, start=1):
        row = pay_tbl.rows[r_idx].cells
        for i, txt in enumerate([a, b, c_val, d]):
            bold = (r_idx == 3 and i in (0, 2, 3))
            add_run(row[i].paragraphs[0], txt, bold=bold, size=10, color=INK)
    small(doc, '', space_after=4)
    clause(doc, '3.3', 'Facturen worden voldaan binnen 14 dagen na factuurdatum op IBAN NL90 RABO 0172 1492 82 t.n.v. Viral Conversions, onder vermelding van het factuurnummer.')
    clause(doc, '3.4', 'Bij overschrijding van de betaaltermijn is de opdrachtgever van rechtswege in verzuim en is de wettelijke (handels)rente verschuldigd. Buitengerechtelijke incassokosten komen voor rekening van de opdrachtgever conform de Wet normering buitengerechtelijke incassokosten.')
    clause(doc, '3.5', 'Eventuele klachten met betrekking tot een factuur worden binnen 14 dagen na factuurdatum schriftelijk en onderbouwd kenbaar gemaakt. Klachten schorten de betalingsverplichting niet op.')

    # Article 4
    article_heading(doc, 4, 'Maandelijkse kosten — hosting en onderhoud')
    clause(doc, '4.1', 'Na de livegang van de webshop worden de maandelijkse kosten voor hosting en onderhoud doorberekend aan de opdrachtgever. Deze kosten staan los van de in artikel 3 genoemde bouwprijs.')
    clause(doc, '4.2', 'Het exacte maandelijkse bedrag wordt nader bepaald op basis van de werkelijke kosten op het moment van livegang (waaronder hostingplan, domeinregistratie, SSL, beveiligings- en monitoringtools, en periodieke updates).')
    clause(doc, '4.3', 'De opdrachtnemer mag het maandelijkse bedrag aanpassen indien de werkelijke kosten daartoe aanleiding geven. Een wijziging wordt minimaal 30 dagen vooraf schriftelijk aangekondigd. De opdrachtgever heeft het recht om bij een ongewenste verhoging de hosting binnen die termijn op te zeggen en de webshop te (laten) migreren.')
    clause(doc, '4.4', 'Indien de opdrachtgever de maandelijkse kosten niet tijdig voldoet, is de opdrachtnemer gerechtigd de hosting en daarmee de online beschikbaarheid van de webshop op te schorten totdat de openstaande bedragen volledig zijn voldaan. Opschorting laat de betalingsverplichting onverlet.')

    # Article 5
    article_heading(doc, 5, 'Uitvoering en planning')
    clause(doc, '5.1', 'De opdrachtnemer start met de werkzaamheden zodra de aanbetaling als bedoeld in artikel 3.2 is bijgeschreven op de in artikel 3.3 genoemde rekening.')
    clause(doc, '5.2', 'Tussen partijen geldt geen fatale opleverdatum. De opdrachtnemer spant zich in om de webshop binnen een redelijke termijn op te leveren en zal de opdrachtgever bij relevante mijlpalen op de hoogte houden.')
    clause(doc, '5.3', 'De verplichting van de opdrachtnemer betreft een inspanningsverplichting en geen resultaatsverplichting, behoudens voor zover een resultaat expliciet uit deze overeenkomst voortvloeit.')
    clause(doc, '5.4', 'Indien de oplevering wordt vertraagd door het uitblijven of onvolledig zijn van door de opdrachtgever aan te leveren content, beslissingen of medewerking, wordt de planning navenant verschoven. De opdrachtnemer is in dat geval niet aansprakelijk voor de daaruit voortvloeiende vertraging.')
    clause(doc, '5.5', 'De opdrachtgever zorgt voor tijdige en volledige aanlevering van onder meer: producttekst en -afbeeldingen, prijzen, voorraadgegevens, KvK- en btw-nummer, logo, huisstijlelementen, juridische teksten (algemene voorwaarden en privacyverklaring) en gegevens voor het betaalproviders-account.')

    # Article 6
    article_heading(doc, 6, 'Revisies en meerwerk')
    clause(doc, '6.1', 'Revisies binnen de in artikel 2.2 omschreven scope zijn inbegrepen in de bouwprijs en worden uitgevoerd tot het moment van goedkeuring door de opdrachtgever, met dien verstande dat partijen zich daarbij redelijk en constructief opstellen.')
    clause(doc, '6.2', 'Werkzaamheden en wijzigingen buiten de scope, uitgevoerd tijdens de bouwfase en vóór de slotbetaling als bedoeld in artikel 3.2, worden aangemerkt als meerwerk. Meerwerk wordt vooraf in overleg vastgesteld en uitgevoerd tegen een uurtarief van € 75,00 per uur exclusief btw, dan wel tegen een vooraf overeengekomen vaste prijs.')
    clause(doc, '6.3', 'Meerwerk wordt achteraf gefactureerd, met daarbij een korte specificatie van de uitgevoerde werkzaamheden en de bestede tijd.')
    clause(doc, '6.4', 'Goedkeuring van de opgeleverde webshop, en daarmee de betaling van de slottermijn als bedoeld in artikel 3.2, mag door de opdrachtgever niet op onredelijke gronden worden uitgesteld of onthouden. Indien de opdrachtgever binnen 14 dagen na schriftelijke oplevering geen onderbouwde gebreken meldt, geldt de webshop als goedgekeurd.')
    clause(doc, '6.5', 'Na oplevering, goedkeuring en ontvangst van de volledige betaling als bedoeld in artikel 3, geldt voor incidenteel probleem-oplossen en losse supportverzoeken die niet onder regulier onderhoud (artikel 8) vallen, een uurtarief van € 50,00 per uur exclusief btw. Ook deze werkzaamheden worden vooraf in overleg vastgesteld en achteraf gefactureerd met een specificatie van uitgevoerde werkzaamheden en bestede tijd.')

    # Article 7
    article_heading(doc, 7, 'Oplevering en overdracht')
    clause(doc, '7.1', 'Oplevering vindt plaats in een afgesproken staging-omgeving (of demo-omgeving) waarin de opdrachtgever de webshop volledig kan beoordelen, voordat enige overdracht plaatsvindt.')
    clause(doc, '7.2', 'Na schriftelijke goedkeuring door de opdrachtgever én na ontvangst van de volledige betaling als bedoeld in artikel 3, draagt de opdrachtnemer de webshop, het beheerdashboard, alle relevante toegangsgegevens en — voor zover van toepassing — de domeinregistratie en hostingaccounts over op naam van de opdrachtgever.')
    callout(doc, '7.3 — BELANGRIJK. Tot het moment dat de volledige betaling op de in artikel 3.3 genoemde rekening is bijgeschreven, blijft de opdrachtnemer eigenaar van de geleverde bestanden en broncode. Bestanden, broncode en domeinrechten gaan nooit over op de opdrachtgever vóór volledige betaling.')
    clause(doc, '7.4', 'Na overdracht is de opdrachtgever zelf verantwoordelijk voor het zorgvuldig beheer van toegangsgegevens en het maken van back-ups van eigen content. De opdrachtnemer behoudt een eigen kopie van de oplevering voor archief- en supportdoeleinden.')

    # Article 8
    article_heading(doc, 8, 'Hosting, onderhoud en beschikbaarheid')
    clause(doc, '8.1', 'Onderhoud na livegang omvat in elk geval: noodzakelijke beveiligingsupdates, basismonitoring, periodieke back-ups en het oplossen van defecten in de oorspronkelijk opgeleverde scope.')
    clause(doc, '8.2', 'De opdrachtnemer spant zich in voor een hoge beschikbaarheid van de webshop, maar garandeert geen ononderbroken werking. Onderhoudsvensters worden, waar mogelijk, vooraf gecommuniceerd.')
    clause(doc, '8.3', 'Functionele uitbreidingen of structurele wijzigingen na oplevering vallen niet onder onderhoud en worden uitgevoerd tegen het post-livegang tarief uit artikel 6.5 (€ 50,00 per uur excl. btw), dan wel tegen een vooraf overeengekomen vaste prijs.')
    clause(doc, '8.4', 'Door derden veroorzaakte storingen (waaronder downtime bij de hostingprovider of betaalprovider) vallen buiten de aansprakelijkheid van de opdrachtnemer; zie artikel 10.')

    # Article 9
    article_heading(doc, 9, 'Intellectueel eigendom en rechten')
    clause(doc, '9.1', 'Alle door de opdrachtnemer ontwikkelde broncode, ontwerpen, templates en andere intellectuele eigendomsrechten met betrekking tot de webshop blijven eigendom van de opdrachtnemer tot het moment dat de opdrachtgever het volledig verschuldigde bedrag uit deze overeenkomst heeft voldaan.')
    clause(doc, '9.2', 'Na volledige betaling verkrijgt de opdrachtgever een niet-exclusief, eeuwigdurend gebruiksrecht op de webshop ten behoeve van de eigen onderneming. Op verzoek werken partijen mee aan een eigendomsoverdracht van de broncode op naam van de opdrachtgever.')
    clause(doc, '9.3', 'Door de opdrachtgever aangeleverde content (productteksten, afbeeldingen, logo’s) blijft eigendom van de opdrachtgever. De opdrachtgever staat ervoor in gerechtigd te zijn deze content te gebruiken op de webshop en vrijwaart de opdrachtnemer voor aanspraken van derden ter zake.')
    clause(doc, '9.4', 'De opdrachtnemer mag de webshop in algemene termen vermelden als referentieproject (bijvoorbeeld in een portfolio), tenzij de opdrachtgever daartegen schriftelijk bezwaar maakt.')

    # Article 10
    article_heading(doc, 10, 'Aansprakelijkheid')
    clause(doc, '10.1', 'De aansprakelijkheid van de opdrachtnemer voor schade die de opdrachtgever lijdt als gevolg van een toerekenbare tekortkoming, onrechtmatige daad of anderszins, is beperkt tot het bedrag dat in de betreffende kwestie door de opdrachtnemer feitelijk is gefactureerd, met een maximum gelijk aan de bouwprijs van € 800,00 (excl. btw).')
    clause(doc, '10.2', 'Aansprakelijkheid voor indirecte schade, gevolgschade, gederfde winst of omzet, gemiste besparingen en schade door bedrijfsstagnatie is uitdrukkelijk uitgesloten.')
    clause(doc, '10.3', 'De opdrachtnemer is niet aansprakelijk voor schade veroorzaakt door of namens derden (zoals hosting- of betaalproviders), door onjuist of onvolledig aangeleverde gegevens, of door wijzigingen die de opdrachtgever zelf na oplevering aanbrengt.')
    clause(doc, '10.4', 'De aansprakelijkheidsbeperkingen in dit artikel gelden niet voor opzet of bewuste roekeloosheid van de opdrachtnemer.')
    clause(doc, '10.5', 'Elke aanspraak vervalt door het enkele verloop van twaalf maanden na de gebeurtenis die de schade heeft veroorzaakt.')

    # Article 11
    article_heading(doc, 11, 'Geheimhouding en verwerking persoonsgegevens')
    clause(doc, '11.1', 'Partijen verplichten zich tot geheimhouding van alle vertrouwelijke informatie die zij in het kader van de uitvoering van deze overeenkomst van elkaar verkrijgen.')
    clause(doc, '11.2', 'Voor zover de opdrachtnemer in het kader van het beheer of onderhoud van de webshop persoonsgegevens van klanten van de opdrachtgever verwerkt, geldt de opdrachtnemer als verwerker in de zin van de AVG en sluiten partijen op eerste verzoek een afzonderlijke verwerkersovereenkomst.')
    clause(doc, '11.3', 'De opdrachtgever is en blijft verwerkingsverantwoordelijke voor de via de webshop verzamelde persoonsgegevens en is verantwoordelijk voor het opstellen van de privacyverklaring en het rechtmatig gebruik van deze gegevens.')

    # Article 12
    article_heading(doc, 12, 'Duur en beëindiging')
    clause(doc, '12.1', 'Deze overeenkomst gaat in op de datum van ondertekening en eindigt voor wat betreft de bouwopdracht na oplevering, goedkeuring en volledige betaling als bedoeld in artikel 3 en artikel 7.')
    clause(doc, '12.2', 'Het hosting- en onderhoudsdeel uit artikel 4 wordt aangegaan voor onbepaalde tijd en kan door beide partijen tegen het einde van een kalendermaand met inachtneming van een opzegtermijn van één maand schriftelijk worden opgezegd.')
    clause(doc, '12.3', 'Elke partij is gerechtigd de overeenkomst met onmiddellijke ingang en zonder ingebrekestelling te beëindigen indien:')
    bullet(doc, 'de andere partij in staat van faillissement wordt verklaard, surseance van betaling aanvraagt, of feitelijk haar onderneming staakt;')
    bullet(doc, 'de andere partij toerekenbaar tekortschiet in de nakoming van een wezenlijke verplichting en die tekortkoming, na schriftelijke ingebrekestelling met een redelijke termijn, niet herstelt.')
    clause(doc, '12.4', 'Bij beëindiging vóór oplevering wegens een aan de opdrachtgever toerekenbare reden, blijft de aanbetaling als bedoeld in artikel 3.2 verschuldigd en worden tot dat moment verrichte werkzaamheden gefactureerd op basis van het uurtarief uit artikel 6.2.')
    clause(doc, '12.5', 'Bij beëindiging vóór oplevering wegens een aan de opdrachtnemer toerekenbare reden, worden tot dat moment niet uitgevoerde werkzaamheden niet in rekening gebracht en wordt het reeds aanbetaalde bedrag verrekend met de tot dan toe daadwerkelijk verrichte werkzaamheden.')

    # Article 13
    article_heading(doc, 13, 'Toepasselijk recht en geschillen')
    clause(doc, '13.1', 'Op deze overeenkomst en alle daaruit voortvloeiende of daarmee samenhangende rechtsverhoudingen is uitsluitend Nederlands recht van toepassing.')
    clause(doc, '13.2', 'Geschillen die voortvloeien uit of samenhangen met deze overeenkomst worden in eerste instantie via overleg tussen partijen opgelost. Indien partijen er niet uitkomen, worden geschillen voorgelegd aan de bevoegde rechter in het arrondissement Rotterdam.')
    clause(doc, '13.3', 'Wijzigingen op of aanvullingen van deze overeenkomst zijn slechts geldig indien schriftelijk overeengekomen tussen partijen.')
    clause(doc, '13.4', 'Indien één of meer bepalingen van deze overeenkomst geheel of gedeeltelijk nietig of vernietigbaar blijken, blijven de overige bepalingen onverkort van kracht. Partijen vervangen de nietige of vernietigde bepaling door een geldige bepaling die zo veel mogelijk aansluit bij de strekking van de oorspronkelijke bepaling.')

    signature_block(doc, client)
    small(doc, 'Viral Conversions · KvK 99922533 · NL267949303B01', size=8.5, color=MUTED, space_after=0)
    doc.save(output_path)


# ── HTML renderer — reads the template and swaps tokens ──────────────────
def render_contract_html(client, docx_url=''):
    """Read the canonical HTML template and substitute CLIENT_* tokens.
    `client` is a dict of placeholder→value. `docx_url` overrides the
    download-bar href so the per-client URL works on a Flask route."""
    tpl_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Contract_DeCargoWinkel.html')
    with open(tpl_path, 'r', encoding='utf-8') as f:
        html = f.read()
    fallback_docx = docx_url or 'Contract_DeCargoWinkel.docx'
    tokens = dict(client)
    tokens.setdefault('CONTRACT_DATUM', '__________________')
    tokens['DOCX_URL'] = fallback_docx
    for key, val in tokens.items():
        html = html.replace('{{' + key + '}}', str(val or '__________________'))
    return html


def render_contract_html_for_onboarding_row(row, docx_url=''):
    return render_contract_html(_from_onboarding_row(row), docx_url=docx_url)


def build_contract_docx_for_onboarding_row(row, output_path):
    return build_contract_docx(_from_onboarding_row(row), output_path)


# ── CLI: regenerate the static De Cargo Winkel DOCX fallback ─────────────
if __name__ == '__main__':
    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Contract_DeCargoWinkel.docx')
    build_contract_docx(DEFAULT_CLIENT, out_path)
    print(f'Saved: {out_path}')
